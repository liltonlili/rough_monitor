#coding:utf8
import requests
import time as tt
from docx import Document
from docx.shared import Inches
from lxml import etree
import os
import sys
import logging
import logging.config
reload(sys)
sys.setdefaultencoding("utf8")


def get_page(content):
    tpage = "0"
    try:
        root = etree.HTML(content)
        tpage = root.xpath('/html/body/./div[@id="wrap_container"]/table/tr[1]/td/table/tr/td[2]/text()[2]')
        tpage = tpage[0].replace("\n","").replace("\t","").replace("\r","").replace("|","").split("/")[1]
    except Exception,err:
        logging.getLogger().error("Can't extract page, err:%s" %err)
        print ("can't extrace page")
    return int(tpage)

data = {
    'pwdlevel' : 'Y',
    'loginType' : 1,
    'userName' : 'liltonlili',
    'pwd' : 'taoguba366381',
    'save' : 'Y'
}


general_headers = {
    "Accept":"text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8",
    "Accept-Encoding":"gzip, deflate, sdch",
    "Accept-Language":"zh-CN,zh;q=0.8",
    "Connection":"keep-alive",
    "Host":"www.taoguba.com.cn",
    "Upgrade-Insecure-Requests":"1",
    "User-Agent":"Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/51.0.2704.103 Safari/537.36"
}


headers = {
    "Accept":"text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8",
    "Accept-Encoding":"gzip, deflate",
    "Accept-Language":"zh-CN,zh;q=0.8",
    "Cache-Control":"max-age=0",
    "Connection":"keep-alive",
    "Content-Type":"application/x-www-form-urlencoded",
    "Host":"www.taoguba.com.cn",
    "Origin":"http://www.taoguba.com.cn",
    "Referer":"http://www.taoguba.com.cn/gotoLogin",
    "Upgrade-Insecure-Requests":1,
    "User-Agent":"Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/51.0.2704.103 Safari/537.36"
}


# general_headers = {
#     "Accept":"text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8",
# "Accept-Encoding":"gzip, deflate, sdch",
# "Accept-Language":"zh-CN,zh;q=0.8",
# "Cache-Control":"max-age=0",
# "Connection":"keep-alive",
# "Host":"www.taoguba.com.cn",
# "Upgrade-Insecure-Requests":"1",
# "User-Agent":"Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/51.0.2704.103 Safari/537.36"
# }

image_headers = {
    "Accept":"text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8",
    "Accept-Encoding":"gzip, deflate, sdch",
    "Accept-Language":"zh-CN,zh;q=0.8",
    "Cache-Control":"max-age=0",
    "Connection":"keep-alive",
    "Host":"image.taoguba.com.cn",
    "Upgrade-Insecure-Requests":1,
    "User-Agent":"Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/51.0.2704.103 Safari/537.36"
}



# 从页面中解析出回复的网址，同时，访问该内容，并抽取其内容
def extract_hrefs(content):
    # with open("./debug.htm",'w') as dFileHandler:
    #     dFileHandler.write(content)
    root = etree.HTML(content)
    hrefs = root.xpath('/html/body/./div[@id="wrap_container"]/table/tr[2]/td//table/tr//td[@class="table_bottom01"]/parent::*//td[@align="left"]/a/@href')
    hrefs.reverse()
    return hrefs


def get_detailed_content(href):

    global pic_count, docHandler, mydir, txtHandler
    # #for debug
    # href = 'http://www.taoguba.com.cn/Reply/1149649/20686733#20686733'
    print href
    logging.getLogger().info("detail href:%s" %href)
    try:
        hrefId = "reply"+href.split("#")[1]
        hrefId = hrefId.strip()
        detail_count = 5
        while detail_count > 0:
            try:
                r = s.get(href,headers = general_headers)
                if r.status_code == 200:
                    break
                tt.sleep(10)
            except:
                pass
            finally:
                detail_count -= 1
        content = r.content.decode('utf8').replace(u'\x00',u'')
        root = etree.HTML(content)

        count = 0
        try:
            name = root.xpath('//div[@class="pcnr_wz"]/p[@id="%s"]/parent::*/parent::*//div[@class="left pcyc_l"]/a[2]/text()'%hrefId)[0]
            time = root.xpath('//div[@class="pcnr_wz"]/p[@id="%s"]/parent::*/parent::*//div[@class="left pcyc_l"]/span/text()'%hrefId)[0]

            # 正文内容
            mainContent = root.xpath('//div[@class="pcnr_wz"]/p[@id="%s"]'%hrefId)[0].xpath("string(.)")

            # mainContent = root.xpath('//div[@class="pcnr_wz"]/p[@id="%s"]/text()'%hrefId)
            # #关键词
            # keywords = root.xpath('//div[@class="pcnr_wz"]/p[@id="%s"]/a/text()'%hrefId)
            #
            # ## 图片后的文字
            # tailContent = root.xpath('//div[@class="pcnr_wz"]/p[@id="%s"]/parent::*/text()'%hrefId)
        except:
            name = root.xpath('//div[@class="p_wenz"]//div[@class="left p_tationl"]/a[1]/span/text()')[0]
            time = root.xpath('//div[@class="p_wenz"]//div[@class="left p_tationl"]/span[1]/text()')[0]

            # 正文内容
            mainContent = root.xpath('//div[@class="p_coten"]')[0].xpath("string(.)")
            # # p/text为正文内容，可能会被关键词打断
            # mainContent = root.xpath('//div[@class="p_coten"]/text()')
            #
            # #关键词
            # keywords = root.xpath('//div[@class="p_coten"]/a/text()')
            #
            # ## 图片后的文字
            # tailContent = root.xpath('//div[@class="p_coten"]/div[@align="center"]/text()')
        time_date = time[:10]
        prefix = u"\n----%s, %s----:" %(name,time)

        #图片
        try:
            pics = root.xpath('//div[@class="pcnr_wz"]/p[@id="%s"]/parent::*/div[@align="center"]/img/@data-original'%hrefId)
        except:
            pics = root.xpath('//div[@class="p_coten"]/div[@align="center"]/img/@data-original')

        try:
            txtHandler.write(prefix.encode('utf8'))
        except:
            try:
                txtHandler.write(prefix)
            except:
                try:
                    txtHandler.write(prefix.encode('gbk'))
                except:
                    try:
                        txtHandler.write(prefix.decode('utf8'))
                    except:
                        txtHandler.write(prefix.decode('gbk'))
        docHandler.add_paragraph("%s\n\n"%prefix)

        for pic in pics:
            if len(pic) <= 5:
                continue
            if '.jpg' in pic:
                file_name = "./%s_%s.jpg"%(time_date,pic_count)
            elif ".png" in pic:
                file_name = "./%s_%s.png"%(time_date,pic_count)
            elif ".bmp" in pic:
                file_name = "./%s_%s.bmp"%(time_date,pic_count)
            else:
                file_name = "./%s_%s.jpg"%(time_date,pic_count)
            file_name = os.path.join(mydir,"pic/%s"%file_name)
            try:
                with open(file_name, 'wb') as fhandler:
                    rPic = s.get(url = pic, headers = image_headers)
                    fhandler.write(rPic.content)
                docHandler.add_picture(file_name,width=Inches(6))
                try:
                    txtHandler.write(file_name)
                except:
                    txtHandler.write(file_name.encode('utf8'))
            except Exception,err:
                logging.getLogger().error("Error when add picture, err:%s, file:%s" %(err, file_name))
                # print err
                # print file_name
            pic_count += 1


        mainContent = mainContent.replace(" ","").replace(u'\xa0','\n').replace("\t","").replace("\n\n","\n").replace("\n\n","\n").replace("\n\n","\n").replace("\r\n","")
        mainContent = mainContent.split(u"[淘股吧]")[0]
        #引用
        try:
            referContent = root.xpath('//div[@class="pcnr_wz"]/p[@id="%s"]/span'%hrefId)[1].xpath('string(.)').replace(" ","").replace(' ','')
            referContent = referContent.replace(u"发表",u"发表\n\t")
            referContent = referContent.replace(u'\xa0','').replace(" ","").strip()
        except:
            referContent = 'None'

        mainContents = "\t"+mainContent
        textContent = "%s\n\n##%s" %(mainContents, referContent)
        try:
            txtHandler.write(textContent.encode('utf8'))
        except:
            try:
                txtHandler.write(textContent)
            except:
                try:
                    txtHandler.write(textContent.encode('gbk'))
                except:
                    try:
                        txtHandler.write(textContent.decode('utf8'))
                    except:
                        txtHandler.write(textContent.decode('gbk'))
        try:
            docHandler.add_paragraph(textContent)
        except:
            try:
                docHandler.add_paragraph(textContent.decode("utf8"))
            except:
                docHandler.add_paragraph(textContent.decode("gbk"))
        tt.sleep(8)
    except:
        docHandler.add_paragraph(href)
        txtHandler.write(href)

if __name__ == "__main__":
    global s, pic_count, docHandler, mydir

    # logging.basicConfig(level=logging.WARNING,
    #                 format='%(asctime)s %(filename)s[line:%(lineno)d] %(levelname)s %(message)s',
    #                 datefmt='%a, %d %b %Y %H:%M:%S',
    #                 filename=os.path.join(u"D:/Money/lilton_code/Market_Mode/learnModule/令胡冲/log","log.txt"),
    #                 filemode='w')

    logging.config.fileConfig("./conf/conf_log.txt", defaults={'logdir': os.path.join(u"D:/Money/lilton_code/Market_Mode/learnModule/logs/fqza","")})
    logging.getLogger().info("report main startup")

    ID = '190857' #
    logging.getLogger().info(ID)
    mydir = u'D:/Money/lilton_code/Market_Mode/learnModule/fqza'
    global txtHandler
    txtHandler = open(os.path.join(mydir,"record.txt"),'w')
    docHandler = Document()

    s = requests.Session()
    login_url = 'http://www.taoguba.com.cn/newLogin'
    pic_count = 0
    r = s.post(url=login_url, headers=headers, data=data, allow_redirects=True)


    index_url = 'http://www.taoguba.com.cn/index'
    r = s.get(index_url,headers = general_headers )
    user_url = 'http://www.taoguba.com.cn/moreReply?userID=%s'%ID

    #主页
    r1 = s.get(user_url,headers = general_headers )
    tpages = int(get_page(r1.content.decode("utf8")))
    # hrefs = extract_hrefs(r1.content.decode("utf8"))
    #得到每个回复的href之后，进行detail的解析
    # for href in hrefs:
    #     href = "http://www.taoguba.com.cn/" + href
    #     get_detailed_content(href)
    try:
        for page in range(tpages,0,-1):
            logging.getLogger().info("page:%s" %page)
            url = 'http://www.taoguba.com.cn/moreReply?pageNum=%s&pageNo=%s&userID=%s' %(tpages, page, ID)
            logging.getLogger().info("url:%s" %url)
            page_count = 20
            status = False
            try:
                while page_count > 0:
                    try:
                        r_tmp = s.get(url,headers = general_headers )
                        if r_tmp.status_code == 200:
                            status = True
                            break
                    except:
                        pass
                    tt.sleep(10)
                    page_count -= 1
                if not status:
                    logging.getLogger().error("page %s failed, will skip" %page)
                    continue
                hrefs = extract_hrefs(r_tmp.content.decode("utf8"))
                #得到每个回复的href之后，进行detail的解析
                for href in hrefs:
                    # #debug
                    # href = "http://www.taoguba.com.cn/Reply/1149649/17971214#17971214"
                    href = "http://www.taoguba.com.cn/" + href
                    get_detailed_content(href)
                    # break
            except Exception, err:
                logging.getLogger().error("Encount err when handle page:%s, err:%s" %(page, err))
                print "fail at page:%s" %page
            # break
    except:
        logging.getLogger().exception("Unexpected err:%s"%err)
        # pass
    finally:
        txtHandler.close()
        docHandler.save(os.path.join(mydir,"./%s.docx"%ID))